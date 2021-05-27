import logging
import re
import os
from bs4 import BeautifulSoup
from django.conf import settings

from fpdf import FPDF
from datetime import datetime,date
from openpyxl import load_workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from docx import Document
from docx.shared import Pt

class CustomPDF(FPDF):

    def header(self):
        # Set up a logo
        self.image('./static/img/juniper-networks-black-rgb.png',12,20,144)
        self.set_font('Times', '', 10)

        # Add an address
        self.cell(100)
        self.cell(0, 10, '2251 Corporate Park Dr #100', 0, 0, 'R')
        self.ln(10)

        self.cell(100)
        self.cell(0, 10, 'Herndon, VA 20171', 0, 0, 'R')
        self.ln(10)

        self.cell(100)
        self.cell(0, 10, 'Phone: (571) 203-1700', 0, 0, 'R')
        # Line break
        self.ln(10)
        # sub_title = first_line + ' - ' + last_line
        # self.cell(0, 5, sub_title, 0, 0, 'C')

    def footer(self):
        self.set_y(-10)

        self.set_font('Arial', 'I', 8)
        """Add date report was run"""
        report_date = str(datetime.today())
        self.cell(0, 10, report_date, 0, 0, 'L')
        # Add a page number
        page = 'Page: ' + str(self.page_no()) + '/{nb}'
        self.cell(0, 10, page, 0, 0, 'R')

def process_upload_log(filename,startdate,enddate):
    with open(filename, 'r') as upload_log:
        logs = []
        for line in upload_log:
            line = re.sub('\s+', ',', line)
            line = line.strip().split(',')
            if startdate <= datetime.strptime(line[0],'%Y-%m-%d').date() <= enddate:
                if line[13]=="|uploaded_file:":
                    newlog={}
                    newlog['uploaddate']=line[0]
                    newlog['uploadtime']=line[1]
                    newlog['uploaduser']=line[9]
                    newlog['uploadip']=line[12]
                    newlog['filename']=line[14]
                    newlog['deleted']='N'
                    logs.append(newlog)
                elif line[13]=="|deleted_file:":
                    newlog={}
                    newlog['uploaddate']=line[0]
                    newlog['uploadtime']=line[1]
                    newlog['uploaduser']=line[9]
                    newlog['uploadip']=line[12]
                    newlog['filename']=line[14]
                    newlog['deleted']='Y'
                    logs.append(newlog)
                elif line[13]=="|verified_file:":
                    for logentry in logs:
                        if logentry['filename']==line[14] and logentry['deleted']=='N':
                            logentry['verifydate']=line[0]
                            logentry['verifytime']=line[1]
                            logentry['verifyuser']=line[9]
                            logentry['verifyip']=line[12]
                            continue
    return logs

def process_download_log(filename,startdate,enddate):
    with open(filename, 'r') as dload_log:
        logs = []
        for line in dload_log:
            line = re.sub('\s+', ',', line)
            line = line.strip().split(',')
            if startdate <= datetime.strptime(line[0],'%Y-%m-%d').date() <= enddate:
                newlog={}
                newlog['dloaddate']=line[0]
                newlog['dloadtime']=line[1]
                newlog['dloaduser']=line[9]
                newlog['dloadip']=line[12]
                newlog['filename']=line[14]
                logs.append(newlog)
    return logs

def generate_upload_report_pdf(log_file,pdf_file,startdate,enddate):
    pdf = CustomPDF('L','pt','Legal')
    # Create the special value {nb}
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font('Arial', 'B', 12)
    title = 'Upload Report for ' + str(startdate) + ' - ' + str(enddate)
    pdf.cell(0, 12, title, 0, 0, 'C')
    pdf.ln(24)
    logs=process_upload_log(log_file,startdate,enddate)
    if len(logs) == 0:
        pdf.set_font('Arial','B',9)
        pdf.cell(0,12,txt="No uploads in the last week")
    else:
        pdf.set_font('Arial', 'BU', 9)
        line_no = 1

        pdf.cell(0, 10, "Date                   Time                  User                                                      IP                         Deleted?    Verifier                                                     Verify Date       File Uploaded".ljust(300), 0,1)
        pdf.set_font('Courier', '', 10)
        for logentry in logs:
            if 'verifyuser' in logentry:
                pdf.cell(0, 10, txt="{} {}   {}{}{}    {} {} {}".format(logentry['uploaddate'], logentry['uploadtime'], logentry['uploaduser'].ljust(26), logentry['uploadip'].ljust(15), logentry['deleted'],logentry['verifyuser'].ljust(26),logentry['verifydate'],logentry['filename']), ln=1)
            elif logentry['deleted'] == 'Y':
                pdf.cell(0, 10, txt="{} {}   {}{}{}                                          {}".format(logentry['uploaddate'], logentry['uploadtime'], logentry['uploaduser'].ljust(26), logentry['uploadip'].ljust(15),logentry['deleted'],logentry['filename']), ln=1)
            else:
                pdf.cell(0, 10, txt="{} {}   {}{}{}    {}            {}".format(logentry['uploaddate'], logentry['uploadtime'], logentry['uploaduser'].ljust(26), logentry['uploadip'].ljust(15),logentry['deleted'],"unverified".ljust(26),logentry['filename']), ln=1)
            line_no += 1
    pdf.output(pdf_file)
    return len(logs)

def generate_download_report_pdf(log_file,pdf_file,startdate,enddate):
    pdf = CustomPDF('L','pt','Letter')
    # Create the special value {nb}
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font('Arial', 'B', 12)
    title = 'Download Report for ' + str(startdate) + ' - ' + str(enddate)
    pdf.cell(0, 12, title, 0, 0, 'C')
    pdf.ln(24)
    logs=process_download_log(log_file,startdate, enddate)
    if len(logs) == 0:
        pdf.set_font('Arial','B',9)
        pdf.cell(0,12,txt="No downloads in the last week")
    else:
        pdf.set_font('Arial', 'BU', 9)
        line_no = 1

        pdf.cell(0, 10, "Date                   Time                  User                                                       IP                                Filename".ljust(260), 0,1)
        pdf.set_font('Courier', '', 10)
        for logentry in logs:
            pdf.cell(0, 10, txt="{} {}   {}{}{}".format(logentry['dloaddate'], logentry['dloadtime'], logentry['dloaduser'].ljust(26), logentry['dloadip'].ljust(15),logentry['filename']), ln=1)
            line_no += 1
    pdf.output(pdf_file)
    return len(logs)

def generate_upload_report_xlsx(log_file,xlsx_file,startdate,enddate):
    logs=process_upload_log(log_file,startdate, enddate)
    template_file=os.path.join(settings.BASE_DIR,'templates/reports/ul_template.xlsx')
    wb=load_workbook(template_file)
    ws=wb.active
    report_title="Upload Report for " + str(enddate)
    ws.title=report_title
    dltable=ws._tables[0]
    dlstyle=dltable.tableStyleInfo
    for i,dload in enumerate(logs):
        ws.cell(row=3+i,column=2).value=datetime.strptime(dload['uploaddate'],'%Y-%m-%d').date()
        ws.cell(row=3+i,column=2).number_format="YYYY-MM-DD"
        ws.cell(row=3+i,column=3).value=datetime.strptime(dload['uploadtime'],'%H:%M:%S').time()
        ws.cell(row=3+i,column=3).number_format="h:mm:ss"
        ws.cell(row=3+i,column=4).value=dload['uploaduser']
        ws.cell(row=3+i,column=5).value=dload['uploadip']
        ws.cell(row=3+i,column=9).value=dload['filename']
        ws.cell(row=3+i,column=6).value=dload['deleted']
        if 'verifyuser' in dload:
            ws.cell(row=3+i,column=7).value=dload['verifyuser']
            ws.cell(row=3+i,column=8).value=dload['verifyip']
        elif dload['deleted']=='N':
            ws.cell(row=3+i,column=7).value="Unverified"


    dltable.ref='B2:I{}'.format(len(logs)+2)
    ws._tables[0]=dltable
    wb.save(xlsx_file)
    return len(logs)

def generate_download_report_xlsx(log_file,xlsx_file,startdate,enddate):
    logs=process_download_log(log_file,startdate, enddate)
    template_file=os.path.join(settings.BASE_DIR,'templates/reports/dl_template.xlsx')
    wb=load_workbook(template_file)
    ws=wb.active
    report_title="Download Report for " + str(enddate)
    ws.title=report_title
    dltable=ws._tables[0]
    dlstyle=dltable.tableStyleInfo
    for i,dload in enumerate(logs):
        ws.cell(row=3+i,column=2).value=datetime.strptime(dload['dloaddate'],'%Y-%m-%d').date()
        ws.cell(row=3+i,column=2).number_format="YYYY-MM-DD"
        ws.cell(row=3+i,column=3).value=datetime.strptime(dload['dloadtime'],'%H:%M:%S').time()
        ws.cell(row=3+i,column=3).number_format="h:mm:ss"
        ws.cell(row=3+i,column=4).value=dload['dloaduser']
        ws.cell(row=3+i,column=5).value=dload['dloadip']
        ws.cell(row=3+i,column=6).value=dload['filename']
    dltable.ref='B2:F{}'.format(len(logs)+2)
    ws._tables[0]=dltable
    wb.save(xlsx_file)
    return len(logs)

def generate_upload_report_docx(log_file,docx_file,startdate,enddate):
    logs=process_upload_log(log_file,startdate, enddate)
    template_file=os.path.join(settings.BASE_DIR,'templates/reports/ul_template.docx')
    report_title="Upload Report for " + str(startdate) + " - " + str(enddate)
    dlreport=Document(template_file)
    for paragraph in dlreport.paragraphs:
        if paragraph.text == "Upload Report":
            paragraph.text=report_title
            titlefont=paragraph.runs[0].font
            titlefont.name = 'Calibri'
            titlefont.size=Pt(16)
            titlefont.bold= True


    if len(logs) > 0:
        dltable=dlreport.tables[0]
        tablefont=dltable.style.font
        tablefont.name='Courier'
        tablefont.size=Pt(8)
        for dload in logs:
            row_cells=dltable.add_row().cells
            row_cells[0].text=dload['uploaddate']
            row_cells[1].text=dload['uploadtime']
            row_cells[2].text=dload['uploaduser']
            row_cells[3].text=dload['uploadip']
            row_cells[4].text=dload['deleted']
            if 'verifyuser' in dload:
                row_cells[5].text=dload['verifyuser']
                row_cells[6].text=dload['verifyip']
            elif dload['deleted']=='N':
                row_cells[5].text="Unverified"
            row_cells[7].text=dload['filename']
    else:
        dlreport.add_paragraph("No uploads for this period")
    dlreport.save(docx_file)
    return len(logs)

def generate_download_report_docx(log_file,docx_file,startdate,enddate):
    logs=process_download_log(log_file,startdate, enddate)
    template_file=os.path.join(settings.BASE_DIR,'templates/reports/dl_template.docx')
    report_title="Download Report for " + str(startdate) + " - " + str(enddate)
    dlreport=Document(template_file)
    for paragraph in dlreport.paragraphs:
        if paragraph.text == "Weekly Download Report":
            paragraph.text=report_title
            titlefont=paragraph.runs[0].font
            titlefont.name = 'Calibri'
            titlefont.size=Pt(16)
            titlefont.bold= True

    if len(logs) > 0:

        dltable=dlreport.tables[0]
        tablefont=dltable.style.font
        tablefont.name='Courier'
        tablefont.size=Pt(8)
        for dload in logs:
            row_cells=dltable.add_row().cells
            row_cells[0].text=dload['dloaddate']
            row_cells[1].text=dload['dloadtime']
            row_cells[2].text=dload['dloaduser']
            row_cells[3].text=dload['dloadip']
            row_cells[4].text=dload['filename']
    else:
        dlreport.add_paragraph("No downloads for this period")
    dlreport.save(docx_file)
    return len(logs)


def do_user_logging(request, action='query'):
    """
    Write page user activity info to log.
    :param request: The request posted to the page.
    :param action:
    the action being logged

    :return:
    """
    logger = logging.getLogger('project_user')
    request_data = ''
    for k, v in request.POST.lists():
        if k not in ('csrfmiddlewaretoken', 'submit', 'submit_records', 'searchBtn', 'sid'):
            request_data = request_data + k + ':'
            for x in v:
                request_data = request_data + x + '|'
    referer = request.META.get('HTTP_REFERER') or ''
    path = request.path or ''
    method = request.method or ''
    logger.info("|action: " + action + "|" + "method: " + method + "|" + "path: " + path + "|" + "http_referer: " + referer + "|" + "userid: " + str(request.user) + "|" + request_data)
    # Track the pages that the user visits
    if 's_pages_visited' in request.session:
        all_pages = request.session['s_pages_visited']
        if path not in request.session['s_pages_visited']:
            all_pages = all_pages + path + '|'
            request.session['s_pages_visited'] = all_pages
    else:
        request.session['s_pages_visited'] = str(str(request.user)) + ':' + path + '|'
    logger.info('|Pages visited by: ' + str(request.session['s_pages_visited']))


def convert_html_to_plain_text(html_email):
    soup = BeautifulSoup(html_email, "html.parser")
    for data in soup(['style', 'script']):
        data.decompose()
    return '\n'.join(soup.stripped_strings)
